#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from flask import Flask, request, jsonify, send_file, render_template_string
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from datetime import date
import copy, os, io, base64
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app = Flask(__name__)

@app.after_request
def add_cors(response):
    response.headers["Access-Control-Allow-Origin"]  = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return response

@app.route("/generuoti", methods=["OPTIONS"])
def generuoti_options():
    return "", 200

GMAIL_USER     = "grota.laboratorija@gmail.com"
RECIPIENT      = "laboratorija@grota.lt"
CLIENT_ID      = "473866566394-r4gv3tfc7srhdbq7p4f29dqptusc64of.apps.googleusercontent.com"
CLIENT_SECRET  = "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"
REFRESH_TOKEN  = "1//06U-Y-DbgZIKhCgYIARAAGAYSNwF-L9Iryoc4UWRm2PerHcLn3GbCG91S0umPHxeIJkTEkNjjyeIwrsUSuRnAW3hVliR899V7_ds"

# ── Įkelti HTML ───────────────────────────────────────────
HTML_FILE = os.path.join(BASE_DIR, "index.html")

@app.route("/")
def index():
    with open(HTML_FILE, encoding="utf-8") as f:
        return f.read()

# ── Word generavimas ──────────────────────────────────────

def generuoti_word(d):
    tipas    = d["tipas"]
    sablonas = os.path.join(BASE_DIR,
        "sablonas_vanduo.docx" if tipas=="vanduo" else "sablonas_gruntas.docx")

    doc       = Document(sablonas)
    paras     = doc.paragraphs
    rodikliai = d["rodikliai"]
    eminiai   = d["eminiai"]
    TNR       = "Times New Roman"

    def _tnr(run, bold=None, size_pt=None):
        run.font.name = TNR
        rPr = run._r.get_or_add_rPr()
        rf  = rPr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
        for a in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'): rf.set(qn(a), TNR)
        if bold is not None: run.font.bold = bold
        if size_pt: run.font.size = Pt(size_pt)

    def set_cell(cell, text, bold=False, size_pt=9, center=False):
        for p in cell.paragraphs:
            for r in p.runs: r.text = ""
        p = cell.paragraphs[0]
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); _tnr(r, bold=bold, size_pt=size_pt)

    def set_para(para, text, size_pt=10):
        if para.runs:
            para.runs[0].text = text; _tnr(para.runs[0], size_pt=size_pt)
            for r in para.runs[1:]: r.text = ""
        else:
            r = para.add_run(text); _tnr(r, size_pt=size_pt)

    def uzpildyti(tbl, rod_gr, n_fixed=4):
        n_cols = len(tbl.rows[1]._tr.findall(qn('w:tc')))
        for ci in range(n_fixed, n_cols):
            ri  = ci - n_fixed
            txt = rod_gr[ri] if ri < len(rod_gr) else ""
            set_cell(tbl.rows[1].cells[ci], txt, bold=bool(txt), size_pt=8, center=True)

        rows_all = list(tbl.rows)
        sd = rows_all[2]; si = rows_all[-1]
        for row in rows_all[2:]: tbl._tbl.remove(row._tr)

        for em in eminiai:
            new_tr = copy.deepcopy(sd._tr); tbl._tbl.append(new_tr)
            row    = tbl.rows[-1]
            set_cell(row.cells[0], em["pavadinimas"], size_pt=8)
            set_cell(row.cells[1], "",                size_pt=8)
            set_cell(row.cells[2], em["data"],        size_pt=8, center=True)
            set_cell(row.cells[3], em["punktas"],     size_pt=8, center=True)
            for ci in range(n_fixed, n_cols):
                ri  = ci - n_fixed
                val = "1" if ri < len(rod_gr) and em["varneles"].get(rod_gr[ri], False) else ""
                set_cell(row.cells[ci], val, size_pt=8, center=True)

        new_tr = copy.deepcopy(si._tr); tbl._tbl.append(new_tr)
        row    = tbl.rows[-1]
        set_cell(row.cells[0], "IŠ VISO", bold=True, size_pt=8)
        for ci in range(1, n_fixed): set_cell(row.cells[ci], "", size_pt=8)
        for ci in range(n_fixed, n_cols):
            ri = ci - n_fixed
            if ri < len(rod_gr):
                sk = sum(1 for em in eminiai if em["varneles"].get(rod_gr[ri], False))
                set_cell(row.cells[ci], str(sk) if sk else "", bold=True, size_pt=8, center=True)
            else:
                set_cell(row.cells[ci], "", size_pt=8)

    for para in paras:
        if para.text.strip() and para.text.strip()[0].isdigit() and len(para.text.strip())==10:
            set_para(para, str(date.today())); break

    for para in paras:
        tl  = para.text.lower()
        tel = d.get("telefonas","").strip()
        if tl.startswith("užsakovas"):
            txt = f"Užsakovas: {d['imone']}"
            if tel: txt += f" ({tel})"
            set_para(para, txt)
        elif tl.startswith("užsakymo pavadinimas"):
            set_para(para, f"Užsakymo pavadinimas: {d.get('uzsakymas','') or '–'}")
        elif tl.startswith("tiriamasis ėminys"):
            set_para(para, f"Tiriamasis ėminys: {d['tiriamasis']}")
        elif tl.startswith("pastabos"):
            p = d.get("pastabos","").strip()
            set_para(para, f"Pastabos: {p}" if p else "")

    tbl    = doc.tables[0]
    grupes = [rodikliai[i:i+13] for i in range(0, len(rodikliai), 13)]
    uzpildyti(tbl, grupes[0])

    if len(grupes) > 1:
        sab_doc = Document(sablonas)
        tbl_el  = tbl._tbl
        for grupe in grupes[1:]:
            sep    = OxmlElement('w:p'); tbl_el.addnext(sep)
            new_el = copy.deepcopy(sab_doc.tables[0]._tbl); sep.addnext(new_el)
            new_tbl= Table(new_el, doc)
            uzpildyti(new_tbl, grupe)
            tbl_el = new_el

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── El. pašto siuntimas ───────────────────────────────────

def gauti_access_token():
    """Gauna naują access token naudojant refresh token."""
    import urllib.request, urllib.parse, json as _json
    data = urllib.parse.urlencode({
        "client_id":     CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "refresh_token": REFRESH_TOKEN,
        "grant_type":    "refresh_token",
    }).encode()
    req    = urllib.request.Request("https://oauth2.googleapis.com/token", data=data)
    resp   = urllib.request.urlopen(req)
    result = _json.loads(resp.read())
    return result["access_token"]

def siusti_email(imone, word_buf, filename):
    try:
        import urllib.request, urllib.parse, json as _json
        access_token = gauti_access_token()

        msg            = MIMEMultipart()
        msg["From"]    = GMAIL_USER
        msg["To"]      = RECIPIENT
        msg["Subject"] = f"Naujas tyrimu uzsakymas - {imone}"
        msg.attach(MIMEText(
            f"Sveiki,\n\nGautas naujas tyrimu uzsakymas nuo: {imone}\n\nUzsakymo forma prisegta.\n\n- Automatinis pranesimas",
            "plain"))
        part = MIMEBase("application", "octet-stream")
        part.set_payload(word_buf.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
        msg.attach(part)

        # Siusti per Gmail API
        raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
        api_url = f"https://gmail.googleapis.com/gmail/v1/users/me/messages/send"
        req_data = _json.dumps({"raw": raw}).encode()
        req = urllib.request.Request(api_url, data=req_data, headers={
            "Authorization": f"Bearer {access_token}",
            "Content-Type":  "application/json",
        })
        urllib.request.urlopen(req)
        print(f"Laiškas išsiųstas: {imone}")
    except Exception as email_err:
        print(f"El. pasto klaida: {email_err}")

# ── OAuth2 callback ──────────────────────────────────────
@app.route("/oauth2callback")
def oauth2callback():
    import urllib.request, urllib.parse, json as _json
    code = request.args.get("code")
    if not code:
        return "Klaida: kodas nerastas", 400
    try:
        data = urllib.parse.urlencode({
            "code":          code,
            "client_id":     "473866566394-r4gv3tfc7srhdbq7p4f29dqptusc64of.apps.googleusercontent.com",
            "client_secret": "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",
            "redirect_uri":  "https://grota-uzsakymai.onrender.com/oauth2callback",
            "grant_type":    "authorization_code",
        }).encode()
        req  = urllib.request.Request("https://oauth2.googleapis.com/token", data=data)
        resp = urllib.request.urlopen(req)
        result = _json.loads(resp.read())
        refresh_token = result.get("refresh_token","NAV")
        return f"<h2>Refresh Token:</h2><p style='word-break:break-all;font-size:1.2rem'><b>{refresh_token}</b></p><p>Nukopijuokite ir atsiuskite!</p>"
    except Exception as e:
        return f"Klaida: {e}", 500

# ── Generuoti maršrutas ───────────────────────────────────

@app.route("/generuoti", methods=["POST"])
def generuoti():
    try:
        d        = request.json
        imone    = d.get("imone","uzsakymas")
        filename = f"{imone}_{date.today()}.docx"

        word_buf  = generuoti_word(d)
        word_buf2 = generuoti_word(d)
        siusti_email(imone, word_buf2, filename)

        word_buf.seek(0)
        return send_file(
            word_buf,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        import traceback
        return jsonify({"klaida": str(e), "traceback": traceback.format_exc()}), 500

if __name__ == "__main__":
    app.run(debug=True, port=5000)
